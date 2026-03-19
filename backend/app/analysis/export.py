"""Generate feedback export documents (PDF and DOCX).

Per DECISION-010: On-demand generation, no persistence. Only user-visible data
is included — no internal IDs, emails, or timestamps.
"""

from __future__ import annotations

import io
from datetime import datetime, timezone
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

# --- Severity colors ---

SEVERITY_COLORS_PDF = {
    "critical": colors.Color(0.8, 0.1, 0.1),
    "warning": colors.Color(0.85, 0.55, 0.0),
    "note": colors.Color(0.5, 0.5, 0.5),
}

SEVERITY_COLORS_DOCX = {
    "critical": RGBColor(0xCC, 0x19, 0x19),
    "warning": RGBColor(0xD9, 0x8C, 0x00),
    "note": RGBColor(0x80, 0x80, 0x80),
}


def generate_feedback_pdf(
    manuscript_title: str,
    genre: str | None,
    summary: dict[str, Any],
    chapters: list[dict[str, Any]],
    document_summary: dict[str, Any] | None = None,
) -> bytes:
    """Generate a PDF report of manuscript feedback.

    Returns the PDF content as bytes.
    """
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=letter,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ExportTitle",
        parent=styles["Title"],
        fontSize=22,
        spaceAfter=6,
    )
    subtitle_style = ParagraphStyle(
        "ExportSubtitle",
        parent=styles["Normal"],
        fontSize=12,
        textColor=colors.grey,
        spaceAfter=20,
    )
    heading_style = ParagraphStyle(
        "ExportH2",
        parent=styles["Heading2"],
        fontSize=14,
        spaceBefore=16,
        spaceAfter=8,
    )
    normal_style = styles["Normal"]
    small_style = ParagraphStyle(
        "Small",
        parent=styles["Normal"],
        fontSize=8,
        leading=10,
    )

    elements: list = []

    # --- Title page ---
    elements.append(Paragraph(_escape(manuscript_title), title_style))
    meta_parts = []
    if genre:
        meta_parts.append(genre)
    meta_parts.append(f"Generated {datetime.now(timezone.utc).strftime('%B %d, %Y')}")
    elements.append(Paragraph(" | ".join(meta_parts), subtitle_style))

    # Summary stats
    summary_data = [
        ["Chapters Analyzed", "Total Issues", "Critical", "Warnings", "Notes"],
        [
            f"{summary.get('chapters_analyzed', 0)} / {summary.get('chapters_total', 0)}",
            str(summary.get("total_issues", 0)),
            str(summary.get("critical", 0)),
            str(summary.get("warning", 0)),
            str(summary.get("note", 0)),
        ],
    ]
    summary_table = Table(summary_data, colWidths=[1.4 * inch] * 5)
    summary_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#333333")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    elements.append(summary_table)
    elements.append(Spacer(1, 16))

    # --- Nonfiction document summary ---
    if document_summary:
        elements.append(Paragraph("Document Synthesis", heading_style))
        assessment = document_summary.get("overall_assessment", "")
        if assessment:
            elements.append(Paragraph(_escape(assessment), normal_style))
            elements.append(Spacer(1, 8))

        score_fields = [
            ("thesis_clarity_score", "Thesis Clarity"),
            ("argument_coherence", "Argument Coherence"),
            ("evidence_density", "Evidence Density"),
            ("tone_consistency", "Tone Consistency"),
        ]
        score_data = [["Dimension", "Score"]]
        for key, label in score_fields:
            val = document_summary.get(key, "")
            if val:
                score_data.append([label, str(val).replace("_", " ").title()])
        if len(score_data) > 1:
            score_table = Table(score_data, colWidths=[2.5 * inch, 2.5 * inch])
            score_table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#444444")),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                        ("FONTSIZE", (0, 0), (-1, -1), 9),
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                        ("TOPPADDING", (0, 0), (-1, -1), 4),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                    ]
                )
            )
            elements.append(score_table)
            elements.append(Spacer(1, 8))

        strengths = document_summary.get("top_strengths", [])
        if strengths:
            elements.append(
                Paragraph(
                    f"<b>Strengths:</b> {_escape(', '.join(strengths))}", normal_style
                )
            )
        priorities = document_summary.get("top_priorities", [])
        if priorities:
            elements.append(
                Paragraph(
                    f"<b>Priorities:</b> {_escape(', '.join(priorities))}", normal_style
                )
            )
        elements.append(Spacer(1, 12))

    # --- Per-chapter breakdown ---
    for ch in chapters:
        ch_num = ch.get("chapter_number", "?")
        ch_title = ch.get("title") or "Untitled"
        word_count = ch.get("word_count")
        wc_str = f" ({word_count:,} words)" if word_count else ""

        elements.append(
            Paragraph(
                f"Chapter {ch_num}: {_escape(ch_title)}{wc_str}", heading_style
            )
        )

        issues = ch.get("issues", [])
        if not issues:
            elements.append(Paragraph("No issues found.", normal_style))
            elements.append(Spacer(1, 8))
            continue

        # Issues table
        table_data = [["Severity", "Type", "Description", "Suggestion"]]
        for issue in issues:
            sev = issue.get("severity", "note")
            table_data.append(
                [
                    sev.upper(),
                    str(issue.get("type", "")).replace("_", " ").title(),
                    _escape(_truncate(issue.get("description", ""), 200)),
                    _escape(_truncate(issue.get("suggestion", ""), 200)),
                ]
            )

        col_widths = [0.7 * inch, 0.9 * inch, 2.7 * inch, 2.7 * inch]
        issues_table = Table(table_data, colWidths=col_widths, repeatRows=1)

        # Build style commands
        style_cmds = [
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#333333")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cccccc")),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ]

        # Color-code severity cells
        for row_idx in range(1, len(table_data)):
            sev = issues[row_idx - 1].get("severity", "note")
            color = SEVERITY_COLORS_PDF.get(sev, colors.grey)
            style_cmds.append(("TEXTCOLOR", (0, row_idx), (0, row_idx), color))
            style_cmds.append(("FONTNAME", (0, row_idx), (0, row_idx), "Helvetica-Bold"))

        issues_table.setStyle(TableStyle(style_cmds))
        elements.append(issues_table)
        elements.append(Spacer(1, 12))

    # Build PDF
    doc.build(elements)
    return buf.getvalue()


def generate_feedback_docx(
    manuscript_title: str,
    genre: str | None,
    summary: dict[str, Any],
    chapters: list[dict[str, Any]],
    document_summary: dict[str, Any] | None = None,
) -> bytes:
    """Generate a DOCX report of manuscript feedback.

    Returns the DOCX content as bytes.
    """
    doc = Document()

    # --- Title ---
    title_para = doc.add_heading(manuscript_title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta_parts = []
    if genre:
        meta_parts.append(genre)
    meta_parts.append(f"Generated {datetime.now(timezone.utc).strftime('%B %d, %Y')}")
    meta_para = doc.add_paragraph(" | ".join(meta_parts))
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in meta_para.runs:
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        run.font.size = Pt(11)

    # --- Summary stats ---
    doc.add_heading("Summary", level=1)
    summary_table = doc.add_table(rows=2, cols=5)
    summary_table.style = "Light Grid Accent 1"
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Analyzed", "Total Issues", "Critical", "Warnings", "Notes"]
    values = [
        f"{summary.get('chapters_analyzed', 0)} / {summary.get('chapters_total', 0)}",
        str(summary.get("total_issues", 0)),
        str(summary.get("critical", 0)),
        str(summary.get("warning", 0)),
        str(summary.get("note", 0)),
    ]
    for i, (h, v) in enumerate(zip(headers, values)):
        summary_table.rows[0].cells[i].text = h
        summary_table.rows[1].cells[i].text = v
        for para in summary_table.rows[0].cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
        for para in summary_table.rows[1].cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Nonfiction document summary ---
    if document_summary:
        doc.add_heading("Document Synthesis", level=1)
        assessment = document_summary.get("overall_assessment", "")
        if assessment:
            doc.add_paragraph(assessment)

        score_fields = [
            ("thesis_clarity_score", "Thesis Clarity"),
            ("argument_coherence", "Argument Coherence"),
            ("evidence_density", "Evidence Density"),
            ("tone_consistency", "Tone Consistency"),
        ]
        scores_with_values = [(label, document_summary.get(key, "")) for key, label in score_fields if document_summary.get(key)]
        if scores_with_values:
            score_table = doc.add_table(rows=1 + len(scores_with_values), cols=2)
            score_table.style = "Light Grid Accent 1"
            score_table.rows[0].cells[0].text = "Dimension"
            score_table.rows[0].cells[1].text = "Score"
            for run in score_table.rows[0].cells[0].paragraphs[0].runs:
                run.bold = True
            for run in score_table.rows[0].cells[1].paragraphs[0].runs:
                run.bold = True
            for idx, (label, val) in enumerate(scores_with_values):
                row = score_table.rows[idx + 1]
                row.cells[0].text = label
                row.cells[1].text = str(val).replace("_", " ").title()

        strengths = document_summary.get("top_strengths", [])
        if strengths:
            p = doc.add_paragraph()
            run = p.add_run("Strengths: ")
            run.bold = True
            p.add_run(", ".join(strengths))

        priorities = document_summary.get("top_priorities", [])
        if priorities:
            p = doc.add_paragraph()
            run = p.add_run("Priorities: ")
            run.bold = True
            p.add_run(", ".join(priorities))

    # --- Per-chapter breakdown ---
    for ch in chapters:
        ch_num = ch.get("chapter_number", "?")
        ch_title = ch.get("title") or "Untitled"
        word_count = ch.get("word_count")
        wc_str = f" ({word_count:,} words)" if word_count else ""

        doc.add_heading(f"Chapter {ch_num}: {ch_title}{wc_str}", level=2)

        issues = ch.get("issues", [])
        if not issues:
            doc.add_paragraph("No issues found.")
            continue

        # Issues table
        table = doc.add_table(rows=1 + len(issues), cols=4)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        # Header row
        for i, header in enumerate(["Severity", "Type", "Description", "Suggestion"]):
            cell = table.rows[0].cells[i]
            cell.text = header
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(9)

        # Data rows
        for row_idx, issue in enumerate(issues):
            row = table.rows[row_idx + 1]
            sev = issue.get("severity", "note")
            sev_cell = row.cells[0]
            sev_cell.text = ""
            sev_run = sev_cell.paragraphs[0].add_run(sev.upper())
            sev_run.bold = True
            sev_run.font.size = Pt(9)
            sev_color = SEVERITY_COLORS_DOCX.get(sev)
            if sev_color:
                sev_run.font.color.rgb = sev_color

            type_text = str(issue.get("type", "")).replace("_", " ").title()
            row.cells[1].text = type_text
            row.cells[2].text = _truncate(issue.get("description", ""), 300)
            row.cells[3].text = _truncate(issue.get("suggestion", ""), 300)

            for cell_idx in range(1, 4):
                for para in row.cells[cell_idx].paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)

    # Write to buffer
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _escape(text: str) -> str:
    """Escape XML special characters for ReportLab Paragraph markup."""
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def _truncate(text: str, max_len: int) -> str:
    """Truncate text to max_len characters, appending ellipsis if needed."""
    if not text:
        return ""
    if len(text) <= max_len:
        return text
    return text[: max_len - 3] + "..."
